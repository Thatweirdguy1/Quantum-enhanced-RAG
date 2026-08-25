r"""Markdown -> DOCX renderer in Amity report format.

Three of the five deliverables are Word documents in the same house style, so the
style lives here once and the prose lives in ``docs/*.md`` where it can be read and
edited as text. The alternative -- 7,000 words of prose embedded in Python string
literals -- makes the content unreviewable and the formatting uncopyable.

House style, matching the department's previous report format:

* Letter paper, 1 inch margins on all four sides
* Times New Roman 12 pt throughout, including headings and table text
* 1.5 line spacing, body paragraphs justified
* Page numbers centred in the footer, restarting at 1 after the title page

Supported Markdown, which is the subset these documents actually use:

===========================  ====================================================
``# ## ### ####``            headings (mapped to Word Heading 1-4 so a TOC field
                             can find them)
``**bold** *italic* `code```  inline runs
``- item`` / ``1. item``     bullet and numbered lists
GFM pipe tables              ``| a | b |`` with a ``|---|---|`` rule
``> quote``                  block quote, indented and italic
``[[TOC]]``                  a real Word TOC field, marked dirty so Word rebuilds
                             it on open
``[[PAGEBREAK]]``            hard page break
``[[CAPTION]] text``         italic centred caption
``---``                      horizontal rule, rendered as a thin bottom border
===========================  ====================================================

YAML-ish front matter between ``---`` fences at the top of the file becomes the
title page. Keys: ``title``, ``subtitle``, ``authors``, ``supervisor``,
``institution``, ``department``, ``degree``, ``date``, ``footer``.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Times New Roman"
SIZE = Pt(12)

# Heading sizes. Kept close to body size on purpose: the department's format uses
# bold and numbering to signal hierarchy rather than large display type.
HEADING_SIZES = {1: Pt(14), 2: Pt(13), 3: Pt(12), 4: Pt(12)}

_INLINE = re.compile(
    r"(\*\*[^*]+?\*\*"          # **bold**
    r"|(?<!\*)\*[^*\n]+?\*(?!\*)"  # *italic*, not part of **
    r"|`[^`\n]+?`)"             # `code`
)


# --------------------------------------------------------------------- low level


def _set_run_font(run, *, bold=False, italic=False, mono=False, size=None) -> None:
    """Apply the house font to one run.

    ``rFonts`` has to be set for ascii, hAnsi and cs separately. Setting only
    ``run.font.name`` leaves the east-asian and complex-script slots on the
    template default, which is how a document ends up looking like Times New Roman
    on one machine and Calibri on another.
    """
    font = run.font
    font.name = "Consolas" if mono else FONT
    font.size = size or SIZE
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(slot), "Consolas" if mono else FONT)


def add_inline(paragraph, text: str, *, bold=False, italic=False, size=None) -> None:
    """Render inline markup into runs on an existing paragraph."""
    for token in _INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, bold=True, italic=italic, size=size)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            # Code is set one point down: Consolas has a larger x-height than
            # Times New Roman and matching the nominal size looks oversized.
            _set_run_font(run, mono=True, italic=italic,
                          size=Pt((size or SIZE).pt - 1))
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, bold=bold, italic=True, size=size)
        else:
            run = paragraph.add_run(token)
            _set_run_font(run, bold=bold, italic=italic, size=size)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    r"""Insert a Word field code, marked dirty so Word recalculates it on open.

    ``w:dirty="true"`` is what makes a TOC populate itself. Without it the reader
    sees an empty heading until they happen to press F9, and a table of contents
    that is blank in the submitted file is worse than no table of contents.
    """
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        run._r.append(OxmlElement("w:t"))
        run._r[-1].text = placeholder
    run._r.append(end)
    _set_run_font(run)


def _bottom_border(paragraph) -> None:
    """A horizontal rule, as a paragraph border rather than a row of hyphens."""
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    ppr.append(borders)


def _restart_page_numbers(section) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), "1")


def _footer_page_number(section, note: str = "") -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.0
    if note:
        add_inline(para, note + "   |   ", size=Pt(10))
    _field(para, "PAGE", "1")


# ------------------------------------------------------------------- doc scaffold


def new_document() -> Document:
    """A blank document with the house style applied to the Normal style."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(slot), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    for level, size in HEADING_SIZES.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = size
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        hr = style.element.get_or_add_rPr()
        hf = hr.find(qn("w:rFonts"))
        if hf is None:
            hf = OxmlElement("w:rFonts")
            hr.append(hf)
        for slot in ("w:ascii", "w:hAnsi", "w:cs"):
            hf.set(qn(slot), FONT)
        style.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, side, Inches(1))
    return doc


def title_page(doc: Document, meta: dict) -> None:
    """The unnumbered front page, followed by a section break.

    Page numbering restarts at 1 on the section after this one, so the title page
    is not counted -- which is what the department's format expects.
    """
    def centred(text, *, size=12, bold=False, italic=False, before=0, after=6):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        para.paragraph_format.line_spacing = 1.5
        if text:
            add_inline(para, text, bold=bold, italic=italic, size=Pt(size))
        return para

    centred("", before=48)
    centred(meta.get("title", ""), size=18, bold=True, after=10)
    if meta.get("subtitle"):
        centred(meta["subtitle"], size=13, italic=True, after=28)
    if meta.get("degree"):
        centred(meta["degree"], size=12, after=30)
    if meta.get("authors"):
        centred("Submitted by", size=12, italic=True, after=6)
        for line in meta["authors"].split(";"):
            centred(line.strip(), size=12, bold=True, after=2)
    if meta.get("supervisor"):
        centred("", after=18)
        centred("Under the guidance of", size=12, italic=True, after=6)
        centred(meta["supervisor"], size=12, bold=True, after=24)
    for key in ("department", "institution", "date"):
        if meta.get(key):
            centred(meta[key], size=12, bold=(key == "institution"), after=6)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Inches(1))
    _restart_page_numbers(section)
    _footer_page_number(section, meta.get("footer", ""))
    # The title page itself gets no footer, so it carries no page number.
    doc.sections[0].footer.is_linked_to_previous = False
    doc.sections[0].footer.paragraphs[0].text = ""


# ---------------------------------------------------------------------- rendering


def _parse_front_matter(text: str) -> tuple[dict, str]:
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    meta: dict[str, str] = {}
    for line in text[3:end].strip().splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            meta[key.strip()] = value.strip()
    return meta, text[end + 4:].lstrip("\n")


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        para = cell.paragraphs[0]
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(2)
        add_inline(para, cell_text, bold=True, size=Pt(11))
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(2)
            # Tables in these documents are often wider than the row they were
            # authored with; a short row pads rather than raising.
            add_inline(para, row[j] if j < len(row) else "", size=Pt(11))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_markdown(doc: Document, text: str) -> None:
    """Render the supported Markdown subset into ``doc``."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "[[PAGEBREAK]]":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if stripped == "[[TOC]]":
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(10)
            _field(para, r'TOC \o "1-3" \h \z \u',
                   "Right-click and choose Update Field to build the contents.")
            i += 1
            continue

        if stripped.startswith("[[CAPTION]]"):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(12)
            add_inline(para, stripped[len("[[CAPTION]]"):].strip(),
                       italic=True, size=Pt(11))
            i += 1
            continue

        if stripped in {"---", "***", "___"}:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(8)
            _bottom_border(para)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            body = stripped[level:].strip()
            para = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            add_inline(para, body, bold=True, size=HEADING_SIZES[min(level, 4)])
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and \
                set(lines[i + 1].strip()) <= set("|-: "):
            rows = [_split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, rows)
            continue

        if stripped.startswith("> "):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.right_indent = Inches(0.4)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(para, " ".join(buf), italic=True)
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        number = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if bullet or number:
            style = "List Number" if number else "List Bullet"
            while i < len(lines):
                s = lines[i].strip()
                m = re.match(r"^(?:[-*]|\d+[.)])\s+(.*)$", s)
                if not m:
                    break
                # A wrapped list item continues on following indented lines.
                body = m.group(1)
                i += 1
                while i < len(lines) and lines[i].startswith("  ") \
                        and lines[i].strip() \
                        and not re.match(r"^(?:[-*]|\d+[.)])\s", lines[i].strip()):
                    body += " " + lines[i].strip()
                    i += 1
                para = doc.add_paragraph(style=style)
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.left_indent = Inches(0.35)
                add_inline(para, body)
            continue

        # Ordinary paragraph: consume until a blank line or a structural marker.
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", ">", "[[", "---")) or \
                    re.match(r"^(?:[-*]|\d+[.)])\s", nxt):
                break
            buf.append(nxt)
            i += 1
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(para, " ".join(buf))


def build(source: Path, target: Path) -> Path:
    """Render one Markdown file to a formatted DOCX."""
    text = source.read_text(encoding="utf8")
    meta, body = _parse_front_matter(text)
    doc = new_document()
    if meta.get("title"):
        title_page(doc, meta)
    else:
        _footer_page_number(doc.sections[0], meta.get("footer", ""))
    render_markdown(doc, body)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target
